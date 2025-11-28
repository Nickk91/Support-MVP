# app/rag/loaders.py - UPDATED with better error handling and fallbacks
from typing import List
import boto3
from botocore.exceptions import ClientError
import tempfile
import os
import logging
from langchain_community.document_loaders import TextLoader, PyPDFLoader
from langchain_core.documents import Document
from dotenv import load_dotenv
import requests

# Load environment variables
load_dotenv()

log = logging.getLogger(__name__)

# Initialize S3 client with explicit credentials
s3_client = boto3.client(
    's3',
    aws_access_key_id=os.getenv('AWS_ACCESS_KEY_ID'),
    aws_secret_access_key=os.getenv('AWS_SECRET_ACCESS_KEY'),
    region_name=os.getenv('AWS_DEFAULT_REGION')
)

def _download_from_s3(s3_url: str) -> str:
    """Download file from S3 to temporary location and return local path"""
    try:
        # Parse S3 URL: https://bucket-name.s3.region.amazonaws.com/key
        if s3_url.startswith('https://'):
            # Extract bucket and key from URL
            parts = s3_url.replace('https://', '').split('.s3.')
            bucket_name = parts[0]
            key = parts[1].split('/', 1)[1] if '/' in parts[1] else parts[1]
            
            log.info(f"🔧 S3 Details - Bucket: {bucket_name}, Key: {key}")
        else:
            # If it's not a full URL, assume it's already a local path
            return s3_url
        
        # Create temporary file
        file_ext = os.path.splitext(key)[1] or '.tmp'
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as temp_file:
            s3_client.download_fileobj(bucket_name, key, temp_file)
            log.info(f"✅ Downloaded from S3: {s3_url} -> {temp_file.name}")
            return temp_file.name
            
    except Exception as e:
        log.error(f"❌ Failed to download from S3 {s3_url}: {e}")
        raise

def extract_s3_key_from_url(s3_url: str) -> str:
    """Extract the S3 key from an S3 URL"""
    if s3_url.startswith('https://'):
        # Format: https://bucket-name.s3.region.amazonaws.com/key
        parts = s3_url.replace('https://', '').split('/')
        if len(parts) > 1:
            return '/'.join(parts[1:])  # Return everything after bucket name
    return s3_url  # Fallback to original

def _load_docx_file(file_path: str) -> List[Document]:
    """Load DOCX files with multiple fallback methods"""
    try:
        # Method 1: Try docx2txt first
        try:
            from langchain_community.document_loaders import Docx2txtLoader
            log.info(f"📝 Loading Word document with Docx2txtLoader: {file_path}")
            return Docx2txtLoader(file_path).load()
        except ImportError:
            log.warning("❌ Docx2txtLoader not available, trying python-docx")
        
        # Method 2: Try python-docx
        try:
            import docx
            doc = docx.Document(file_path)
            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            
            if full_text:
                content = "\n".join(full_text)
                log.info(f"✅ Loaded DOCX with python-docx: {len(content)} characters")
                return [Document(page_content=content, metadata={"source": file_path})]
        except ImportError:
            log.warning("❌ python-docx not available, trying unstructured")
        
        # Method 3: Try unstructured as final fallback
        try:
            from unstructured.partition.docx import partition_docx
            elements = partition_docx(filename=file_path)
            content = "\n".join([str(el) for el in elements])
            log.info(f"✅ Loaded DOCX with unstructured: {len(content)} characters")
            return [Document(page_content=content, metadata={"source": file_path})]
        except ImportError:
            log.error("❌ No DOCX processing libraries available")
        
        # Final fallback: empty document
        log.error("❌ All DOCX loading methods failed")
        return [Document(page_content="", metadata={"source": file_path, "error": "DOCX processing failed"})]
        
    except Exception as e:
        log.error(f"❌ Error loading DOCX file {file_path}: {e}")
        return [Document(page_content="", metadata={"source": file_path, "error": str(e)})]

def _load_local_file(file_path: str) -> List[Document]:
    """Load a local file using appropriate loader with better error handling"""
    p = os.path.abspath(file_path)
    if not os.path.exists(p):
        raise FileNotFoundError(f"Path does not exist: {p}")

    ext = os.path.splitext(p)[1].lower()
    
    try:
        if ext == ".pdf":
            log.info(f"📄 Loading PDF: {p}")
            
            # Method 1: Try PyPDFLoader first
            try:
                loader = PyPDFLoader(p)
                docs = loader.load()
                if docs and any(doc.page_content.strip() for doc in docs):
                    log.info(f"✅ PDF loaded successfully with PyPDF: {len(docs)} pages")
                    return docs
                else:
                    log.warning("⚠️ PyPDF loaded but no content found, trying fallback...")
            except ImportError as e:
                log.warning(f"⚠️ PyPDF not available: {e}, trying fallback...")
            except Exception as e:
                log.warning(f"⚠️ PyPDF failed: {e}, trying fallback...")
            
            # Method 2: Try UnstructuredPDFLoader as fallback
            try:
                from langchain_community.document_loaders import UnstructuredPDFLoader
                loader = UnstructuredPDFLoader(p)
                docs = loader.load()
                if docs and any(doc.page_content.strip() for doc in docs):
                    log.info(f"✅ PDF loaded successfully with UnstructuredPDF: {len(docs)} pages")
                    return docs
                else:
                    log.warning("⚠️ UnstructuredPDF loaded but no content found")
            except ImportError as e:
                log.warning(f"⚠️ UnstructuredPDF not available: {e}")
            except Exception as e:
                log.warning(f"⚠️ UnstructuredPDF failed: {e}")
            
            # Method 3: Try pdfplumber as final fallback
            try:
                import pdfplumber
                content_parts = []
                with pdfplumber.open(p) as pdf:
                    for page_num, page in enumerate(pdf.pages):
                        text = page.extract_text()
                        if text and text.strip():
                            content_parts.append(f"Page {page_num + 1}:\n{text}")
                
                if content_parts:
                    content = "\n\n".join(content_parts)
                    docs = [Document(page_content=content, metadata={"source": p, "total_pages": len(pdf.pages)})]
                    log.info(f"✅ PDF loaded successfully with pdfplumber: {len(pdf.pages)} pages")
                    return docs
            except ImportError as e:
                log.warning(f"⚠️ pdfplumber not available: {e}")
            except Exception as e:
                log.warning(f"⚠️ pdfplumber failed: {e}")
            
            # All methods failed
            log.error(f"❌ All PDF loading methods failed for {p}")
            return [Document(
                page_content="", 
                metadata={"source": p, "error": "PDF processing failed - no content extracted"}
            )]
            
        elif ext in [".docx", ".doc"]:
            return _load_docx_file(p)
            
        else:
            # Handles .txt, .md, .csv, etc.
            log.info(f"📋 Loading text file: {p}")
            try:
                loader = TextLoader(p, encoding="utf-8")
                docs = loader.load()
                if docs and docs[0].page_content.strip():
                    log.info(f"✅ Text file loaded successfully: {len(docs[0].page_content)} characters")
                    return docs
                else:
                    log.warning("⚠️ Text file loaded but no content found")
            except UnicodeDecodeError:
                log.warning(f"UTF-8 failed for {p}, trying latin-1")
                # Fallback to latin-1 encoding if UTF-8 fails
                try:
                    loader = TextLoader(p, encoding="latin-1")
                    docs = loader.load()
                    if docs and docs[0].page_content.strip():
                        log.info(f"✅ Text file loaded with latin-1: {len(docs[0].page_content)} characters")
                        return docs
                    else:
                        log.warning("⚠️ Text file loaded with latin-1 but no content found")
                except Exception as e:
                    log.error(f"❌ Latin-1 encoding also failed: {e}")
            except Exception as e:
                log.error(f"❌ Error loading text file {p}: {e}")
            
            # Return empty document if all text loading methods fail
            return [Document(
                page_content="", 
                metadata={"source": p, "error": "Text processing failed - no content extracted"}
            )]
                
    except Exception as e:
        log.error(f"❌ Error loading file {p}: {e}")
        # Return empty document with error metadata instead of crashing
        return [Document(
            page_content="", 
            metadata={"source": p, "error": str(e)}
        )]
    
def _load_one(path: str) -> List[Document]:
    original_path = path  # Keep original S3 URL
    
    # Check if it's an S3 URL
    if path.startswith('https://') and 'amazonaws.com' in path:
        log.info(f"🔍 Loading from S3: {path}")
        local_path = _download_from_s3(path)
        try:
            docs = _load_local_file(local_path)
            
            # 🎯 CRITICAL: Check if we actually got content
            valid_docs = []
            for doc in docs:
                if doc.page_content and doc.page_content.strip():
                    # Extract S3 key and add to metadata
                    s3_key = extract_s3_key_from_url(path)
                    doc.metadata['original_source'] = original_path
                    doc.metadata['s3_key'] = s3_key
                    doc.metadata['source'] = original_path
                    valid_docs.append(doc)
                else:
                    log.warning(f"⚠️ Empty document content from {path}")
            
            log.info(f"📊 Loaded {len(valid_docs)} valid documents from {path}")
            return valid_docs
            
        finally:
            # Clean up temporary file
            try:
                if os.path.exists(local_path):
                    os.unlink(local_path)
                    log.info(f"🧹 Cleaned up temp file: {local_path}")
            except Exception as e:
                log.warning(f"⚠️ Could not clean up temp file {local_path}: {e}")
    else:
        # Local file path
        docs = _load_local_file(path)
        
        # Filter out empty documents
        valid_docs = []
        for doc in docs:
            if doc.page_content and doc.page_content.strip():
                # For local files, use filename as source
                filename = os.path.basename(path)
                doc.metadata['source'] = filename
                doc.metadata['original_source'] = path
                valid_docs.append(doc)
            else:
                log.warning(f"⚠️ Empty document content from {path}")
        
        log.info(f"📊 Loaded {len(valid_docs)} valid documents from {path}")
        return valid_docs

def load_paths(paths: List[str]) -> List[Document]:
    """Load documents from paths with comprehensive logging"""
    docs: List[Document] = []
    
    log.info(f"🚀 Starting to load {len(paths)} paths")
    
    for raw_path in paths:
        try:
            log.info(f"🔍 Processing file: {raw_path}")
            file_docs = _load_one(raw_path)
            
            if file_docs:
                docs.extend(file_docs)
                log.info(f"✅ Successfully loaded {len(file_docs)} documents from: {raw_path}")
                
                # Log content preview for debugging
                for i, doc in enumerate(file_docs):
                    content_preview = doc.page_content[:100] + "..." if len(doc.page_content) > 100 else doc.page_content
                    log.info(f"   📄 Doc {i}: {len(doc.page_content)} chars, preview: '{content_preview}'")
            else:
                log.warning(f"⚠️ No documents loaded from: {raw_path}")
                
        except Exception as e:
            log.error(f"❌ Failed to load {raw_path}: {e}")
            # Add empty document to continue processing but mark the error
            docs.append(Document(
                page_content="", 
                metadata={"source": raw_path, "error": str(e)}
            ))
            continue
    
    # Final summary
    valid_docs = [doc for doc in docs if doc.page_content and doc.page_content.strip()]
    log.info(f"📊 LOADING SUMMARY: {len(valid_docs)} valid documents out of {len(paths)} paths")
    
    return valid_docs

def load_s3_paths(s3_urls: List[str]) -> List[Document]:
    """Load documents from S3 URLs"""
    s3_client = boto3.client('s3')
    all_docs = []
    
    for s3_url in s3_urls:
        try:
            # Parse bucket and key from S3 URL
            if 's3.amazonaws.com' in s3_url:
                # Format: https://bucket.s3.region.amazonaws.com/key
                bucket_start = s3_url.find('//') + 2
                bucket_end = s3_url.find('.s3')
                bucket = s3_url[bucket_start:bucket_end]
                key_start = s3_url.find(bucket) + len(bucket) + 1
                key = s3_url[key_start:]
            else:
                # Handle other S3 URL formats if needed
                continue
            
            # Download file to temp location
            with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(key)[1]) as tmp_file:
                s3_client.download_fileobj(bucket, key, tmp_file)
                tmp_path = tmp_file.name
            
            # Load using existing loader
            docs = load_paths([tmp_path])
            
            # Update metadata to show S3 source
            for doc in docs:
                doc.metadata['source'] = s3_url
                doc.metadata['s3_bucket'] = bucket
                doc.metadata['s3_key'] = key
            
            all_docs.extend(docs)
            
            # Clean up temp file
            os.unlink(tmp_path)
            
            print(f"✅ Downloaded and processed S3 file: {s3_url}")
            
        except Exception as e:
            print(f"❌ Failed to process S3 URL {s3_url}: {e}")
            continue
    
    return all_docs